from __future__ import annotations

from pathlib import Path


def main() -> int:
    """Generate a Word (.docx) report for the SRM prototype.

    Requires:
        pip install -e ".[report]"
    """

    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except Exception as e:  # pragma: no cover
        print('Missing dependency. Install with: python -m pip install -e ".[report]"')
        print(f"Import error: {e}")
        return 2

    root = Path(__file__).resolve().parents[1]
    out_dir = root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "SRM_Отчет_по_проекту.docx"

    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "ОТЧЁТ ПО ПРОЕКТУ\n"
        "Прототип SRM-системы анализа нарушений закупок\n"
        "на основе онтологической модели правил и комплекса интеллектуальных агентов"
    )
    r.bold = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Учебное заведение: ").bold = True
    meta.add_run("<Название ВУЗа>\n")
    meta.add_run("Факультет/кафедра: ").bold = True
    meta.add_run("<…>\n")
    meta.add_run("Направление: ").bold = True
    meta.add_run("<…>\n")
    meta.add_run("Студент: ").bold = True
    meta.add_run("<ФИО>, группа <…>\n")
    meta.add_run("Руководитель: ").bold = True
    meta.add_run("<ФИО, должность>\n")
    meta.add_run("Год: ").bold = True
    meta.add_run("2026\n")

    doc.add_page_break()

    # Main text
    add_h1(doc, "Аннотация")
    doc.add_paragraph(
        "В рамках работы разработан прототип SRM-системы, предназначенной для "
        "автоматизированного выявления нарушений в данных о закупках. Анализ выполняется на "
        "основе (1) онтологической модели нормативных правил, описываемой в JSON, и "
        "(2) комплекса интеллектуальных агентов-детекторов, каждый из которых отвечает за "
        "отдельный класс нарушений. Прототип поддерживает загрузку данных закупок в формате "
        "JSON и CSV, логирование процесса анализа, формирование отчёта о нарушениях в JSON и "
        "вывод результатов в консоль на русском языке. Дополнительно реализована генерация "
        "графиков (PNG) и опциональный Web API на FastAPI. Для обеспечения качества кода "
        "настроены автоматические проверки и тестирование через GitHub Actions."
    )

    add_h1(doc, "Содержание")
    doc.add_paragraph(
        "1. Введение\n"
        "2. Постановка задачи и требования\n"
        "3. Технологии и инструменты\n"
        "4. Архитектура прототипа\n"
        "5. Модель данных и форматы входных данных\n"
        "6. Онтологическая модель нормативных правил\n"
        "7. Интеллектуальные агенты и логика выявления нарушений\n"
        "8. Отчётность, расчёты, графики и вывод в консоль\n"
        "9. CLI и опциональный Web API\n"
        "10. Тестирование и CI/CD (GitHub Actions)\n"
        "11. Инструкция по запуску\n"
        "12. Заключение и направления развития\n"
        "Приложения"
    )

    add_h1(doc, "1. Введение")
    doc.add_paragraph(
        "Актуальность SRM (Supplier Relationship Management) и систем контроля закупок обусловлена "
        "ростом объёмов данных, сложностью нормативных требований и необходимостью оперативно "
        "выявлять риски и нарушения. Ручная проверка закупочных процедур требует значительных "
        "ресурсов и подвержена человеческим ошибкам. Поэтому востребованы интеллектуальные системы, "
        "которые способны применять формализованные правила и автоматически формировать отчёт о "
        "найденных нарушениях, обеспечивая прозрачность анализа."
    )
    doc.add_paragraph(
        "Целью данного проекта является создание демонстрационного прототипа SRM-системы, который "
        "показывает: (1) как закупочные данные могут быть нормализованы и проанализированы; "
        "(2) как нормативные правила могут быть вынесены в формализованную онтологию; "
        "(3) как мультиагентный подход позволяет расширять систему добавлением новых агентов/правил; "
        "(4) как результаты анализа представляются в отчёте, логах и визуализациях."
    )

    add_h1(doc, "2. Постановка задачи и требования")
    add_h2(doc, "2.1. Требуемый функционал")
    doc.add_paragraph(
        "В прототипе реализованы: загрузка данных о закупках из JSON или CSV; хранение и применение "
        "нормативных правил (онтология в JSON и Python-структуры); интеллектуальный агентный анализ; "
        "выявление нарушений трёх типов (превышение бюджета, недопустимый поставщик, нарушение сроков "
        "поставки); формирование отчёта о нарушениях (JSON) и вывод в консоль на русском языке; "
        "логирование процесса анализа; генерация графиков (PNG) и сохранение в файлы."
    )
    doc.add_paragraph(
        "Дополнительно: предложен интерфейс CLI и опциональный Web API (FastAPI), добавлена обработка "
        "ошибок (например, неверный JSON), настроен CI/CD (GitHub Actions), добавлены docstrings и "
        "типизация."
    )

    add_h2(doc, "2.2. Ограничения и допущения прототипа")
    doc.add_paragraph(
        "Прототип ориентирован на демонстрацию архитектурного подхода, поэтому хранение данных и "
        "результатов реализовано файлово (без БД), онтология представлена в JSON (без RDF/OWL reasoner), "
        "набор нарушений фиксирован, но расширяем по архитектуре."
    )

    add_h1(doc, "3. Технологии и инструменты")
    doc.add_paragraph(
        "Язык: Python 3.10+. Визуализация: matplotlib. Тестирование: pytest. "
        "Статический анализ/форматирование: ruff. CI: GitHub Actions. "
        "Опционально: FastAPI + Uvicorn (Web API)."
    )

    add_h1(doc, "4. Архитектура прототипа")
    doc.add_paragraph(
        "Проект построен модульно: выделены слои данных, онтологии, агентов, логики анализа, отчётности "
        "и интерфейсов."
    )
    doc.add_paragraph(
        "Основные модули:\n"
        "- src/srm/data — загрузка/нормализация данных\n"
        "- src/srm/ontology — онтология правил\n"
        "- src/srm/agents — агенты анализа\n"
        "- src/srm/logic — детекторы, отчёт, графики, консольный вывод\n"
        "- src/srm/cli.py — CLI\n"
        "- src/srm/api/app.py — API (опционально)"
    )
    add_figure_placeholder(
        doc,
        1,
        "Структура проекта в IDE/проводнике",
        "Дерево папок src/srm, examples, tests, .github/workflows и файлы README.md, pyproject.toml.",
        "Показать модульную архитектуру и разделение по слоям.",
    )

    add_h1(doc, "5. Модель данных и форматы входных данных")
    add_h2(doc, "5.1. Нормализованная модель")
    doc.add_paragraph(
        "Нормализованная запись закупки представлена dataclass-моделью ProcurementRecord. "
        "Ключевые поля: идентификатор закупки, поставщик, сумма договора, бюджет, даты договора и "
        "поставки. Дополнительные поля сохраняются в extra для расширяемости."
    )
    add_h2(doc, "5.2. Поддерживаемые форматы JSON/CSV")
    doc.add_paragraph(
        "Поддерживается JSON (массив объектов или объект с ключом procurements) и CSV (таблица). "
        "Загрузка включает проверку формата, преобразования чисел/дат и обработку ошибок."
    )
    add_figure_placeholder(
        doc,
        2,
        "Пример входных данных JSON",
        "Открытый файл examples/procurements.json: поля contract_amount, planned_budget, delivery_due_date, delivery_actual_date.",
        "Показать структуру данных, на которых работает анализ.",
    )

    add_h1(doc, "6. Онтологическая модель нормативных правил")
    doc.add_paragraph(
        "Онтология служит формализованным описанием нормативных правил. Она хранится в отдельном "
        "JSON-файле и преобразуется в типизированные структуры Python (BudgetRule, SupplierRule, DeliveryRule)."
    )
    add_figure_placeholder(
        doc,
        3,
        "Пример онтологии правил (JSON)",
        "Открытый файл examples/rules_ontology.json: секции budget, suppliers, delivery.",
        "Показать, что правила отделены от кода и могут изменяться без переписывания логики.",
    )

    add_h1(doc, "7. Интеллектуальные агенты и логика выявления нарушений")
    doc.add_paragraph(
        "Используется мультиагентная схема: атомарные агенты анализируют данные по своим правилам, "
        "а композитный агент объединяет результаты и устраняет дубли."
    )
    doc.add_paragraph(
        "Реализованные нарушения:\n"
        "1) Превышение бюджета — сравнение contract_amount с бюджетом (planned_budget или бюджет по категории).\n"
        "2) Недопустимый поставщик — проверка whitelist/blacklist из онтологии.\n"
        "3) Нарушение сроков поставки — сравнение due/actual, возможность учитывать недопоставку как нарушение."
    )

    add_h1(doc, "8. Отчётность, расчёты, графики и вывод в консоль")
    doc.add_paragraph(
        "Отчёт формируется в JSON и содержит meta, summary, charts, violations. "
        "В summary рассчитываются агрегаты: суммы, количество нарушений, средняя/максимальная просрочка и т.д."
    )
    add_figure_placeholder(
        doc,
        4,
        "Пример отчёта report.json",
        "Открытый outputs/report.json: раздел summary и несколько элементов violations.",
        "Показать результат работы системы в машиночитаемом виде.",
    )
    doc.add_paragraph(
        "В консоли результаты выводятся на русском языке: итоги и расчёты, ASCII-график по типам нарушений, "
        "список нарушений и пути к файлам результатов."
    )
    add_figure_placeholder(
        doc,
        5,
        "Вывод CLI в консоль",
        "Окно PowerShell после запуска analyze: «Итоги анализа SRM», ASCII-график и список нарушений.",
        "Показать интерактивный результат для пользователя/комиссии.",
    )
    doc.add_paragraph(
        "Графики сохраняются в PNG: распределение нарушений по типам, распределение превышений бюджета и "
        "распределение задержек поставок."
    )
    add_figure_placeholder(
        doc,
        6,
        "Папка outputs/charts с PNG-графиками",
        "Папка outputs/charts: violation_counts.png, budget_overspend_hist.png, delivery_delay_hist.png.",
        "Показать факт генерации визуализаций и состав артефактов.",
    )
    add_figure_placeholder(
        doc,
        7,
        "График «Violations by type»",
        "Открытый outputs/charts/violation_counts.png (столбчатая диаграмма по типам нарушений).",
        "Визуально продемонстрировать распределение нарушений по типам.",
    )
    doc.add_paragraph(
        "Логирование настроено в консоль и файл с ротацией; лог фиксирует этапы анализа и статистику по агентам."
    )
    add_figure_placeholder(
        doc,
        8,
        "Лог-файл srm.log",
        "Открытый outputs/logs/srm.log: записи о запуске анализа и количестве нарушений по агентам.",
        "Показать трассировку выполнения и пригодность для аудита.",
    )

    add_h1(doc, "9. CLI и опциональный Web API")
    doc.add_paragraph(
        "CLI реализует команды analyze (анализ файла закупок и правил) и serve (запуск FastAPI, опционально). "
        "Параметр --open-output открывает отчёт и графики в отдельных окнах (приложения по умолчанию ОС)."
    )
    add_figure_placeholder(
        doc,
        9,
        "Справка CLI (help)",
        "Вывод python -m srm.cli analyze -h (список аргументов и описаний).",
        "Показать пользовательский интерфейс и параметры запуска.",
    )
    doc.add_paragraph(
        "Web API (FastAPI) предоставляет эндпоинты GET /health и POST /analyze (принимает JSON с закупками и онтологией)."
    )
    add_figure_placeholder(
        doc,
        10,
        "Swagger UI FastAPI",
        "Страница http://127.0.0.1:8000/docs: список эндпоинтов и схема запроса/ответа.",
        "Показать готовность к интеграции через HTTP.",
    )

    add_h1(doc, "10. Тестирование и CI/CD (GitHub Actions)")
    doc.add_paragraph(
        "Проект содержит автотесты на pytest и конфигурацию CI на GitHub Actions. "
        "В CI выполняются проверки ruff и запуск тестов."
    )
    add_figure_placeholder(
        doc,
        11,
        "Результат запуска pytest",
        "Окно консоли с выполнением pytest, виден статус passed.",
        "Подтвердить работоспособность ключевых сценариев.",
    )
    add_figure_placeholder(
        doc,
        12,
        "Успешный прогон GitHub Actions",
        "Вкладка Actions на GitHub: зелёный статус workflow, видны шаги lint/test.",
        "Показать CI/CD готовность для комиссии.",
    )

    add_h1(doc, "11. Инструкция по запуску (кратко)")
    doc.add_paragraph(
        "1) Установка:\n"
        "- python -m venv .venv\n"
        "- .\\.venv\\Scripts\\Activate.ps1\n"
        "- python -m pip install -U pip\n"
        '- python -m pip install -e ".[dev]"\n\n'
        "2) Анализ демо-данных:\n"
        "- python -m srm.cli analyze --data .\\examples\\procurements.json --rules .\\examples\\rules_ontology.json "
        "--out .\\outputs\\report.json --charts-dir .\\outputs\\charts --log-dir .\\outputs\\logs\n\n"
        "3) Открыть результаты в отдельных окнах:\n"
        "- добавить --open-output"
    )

    add_h1(doc, "12. Заключение и направления развития")
    doc.add_paragraph(
        "Разработан прототип SRM-системы, демонстрирующий применение онтологии нормативных правил и "
        "мультиагентного анализа для выявления нарушений в закупках. Система поддерживает загрузку данных, "
        "расширяемую архитектуру агентов, генерацию отчётов, логирование и визуализацию, а также имеет CLI "
        "и опциональный Web API. Качество разработки подтверждается тестами и CI pipeline."
    )
    doc.add_paragraph(
        "Направления развития: подключение БД для хранения закупок и истории анализов; расширение онтологии "
        "(формальные связи, RDF/OWL, reasoner); добавление новых агентов (риск-скоринг, рекомендации); "
        "веб-интерфейс (панель мониторинга); интеграция LLM-компонента для объяснений поверх детекторов; "
        "расширение аналитики (тренды, топ-поставщики по нарушениям)."
    )

    add_h1(doc, "Приложение А — Артефакты результатов (пример)")
    doc.add_paragraph(
        "Отчёт: outputs/report.json\nГрафики: outputs/charts/\nЛоги: outputs/logs/srm.log"
    )

    doc.save(out_path)
    print(f"Generated: {out_path}")
    return 0


def add_h1(doc, text: str) -> None:
    doc.add_heading(text, level=1)


def add_h2(doc, text: str) -> None:
    doc.add_heading(text, level=2)


def add_figure_placeholder(
    doc,
    num: int,
    title: str,
    screenshot_description: str,
    purpose: str,
) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Рисунок {num} — {title}").bold = True
    doc.add_paragraph("[МЕСТО ДЛЯ СКРИНШОТА]")
    doc.add_paragraph(f"Описание скриншота: {screenshot_description}")
    doc.add_paragraph(f"Назначение: {purpose}")


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())
